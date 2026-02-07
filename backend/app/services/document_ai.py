"""
Google Document AI 서비스
역할: 문서에서 텍스트 추출 및 OCR 처리

📋 사용 방법
============
1. 서비스 계정 키 파일(.json)을 환경 변수로 설정:
   - GOOGLE_APPLICATION_CREDENTIALS="path/to/key.json"
   
2. 또는 Cloud Run에서는 서비스 계정이 자동 연결됨

3. Document AI 프로세서 생성 필요 (GCP 콘솔에서):
   - 프로젝트: knu-team-03
   - 위치: asia-northeast1 (또는 us)
   - 프로세서 타입: OCR_PROCESSOR 또는 GENERAL_PROCESSOR
"""

import os
from typing import Optional, Dict, Any
from pathlib import Path

# Document AI 클라이언트 (선택적 임포트)
try:
    from google.cloud import documentai
    from google.api_core.client_options import ClientOptions
    DOCUMENT_AI_AVAILABLE = True
except ImportError:
    DOCUMENT_AI_AVAILABLE = False
    documentai = None


class DocumentAIService:
    """
    Google Cloud Document AI 서비스
    
    PDF, 이미지, 스캔 문서에서 고품질 텍스트 추출
    
    Example:
    ```python
    from app.services.document_ai import DocumentAIService
    
    service = DocumentAIService()
    result = await service.process_document(file_path="/tmp/document.pdf")
    print(result["text"])
    ```
    """
    
    def __init__(
        self,
        project_id: str = None,
        location: str = "asia-northeast1",
        processor_id: str = None
    ):
        """
        Document AI 서비스 초기화
        
        Args:
            project_id: GCP 프로젝트 ID (기본값: 환경변수에서 자동 감지)
            location: 프로세서 위치 (asia-northeast1, us, eu)
            processor_id: 프로세서 ID (없으면 기본 OCR 사용)
        """
        self.project_id = project_id or os.getenv("GOOGLE_CLOUD_PROJECT", "knu-team-03")
        self.location = location
        self.processor_id = processor_id or os.getenv("DOCUMENT_AI_PROCESSOR_ID", "")
        
        self._client = None
        self._available = DOCUMENT_AI_AVAILABLE
    
    @property
    def is_available(self) -> bool:
        """Document AI 사용 가능 여부"""
        return self._available and bool(self.processor_id)
    
    def _get_client(self):
        """Document AI 클라이언트 생성 (lazy initialization)"""
        if not self._available:
            raise RuntimeError("google-cloud-documentai 패키지가 설치되지 않았습니다")
        
        if self._client is None:
            # 지역별 엔드포인트 설정
            opts = ClientOptions(
                api_endpoint=f"{self.location}-documentai.googleapis.com"
            )
            self._client = documentai.DocumentProcessorServiceClient(
                client_options=opts
            )
        
        return self._client
    
    def _get_processor_name(self) -> str:
        """프로세서 리소스 이름 생성"""
        return (
            f"projects/{self.project_id}/"
            f"locations/{self.location}/"
            f"processors/{self.processor_id}"
        )
    
    async def process_document(
        self,
        file_path: str = None,
        file_content: bytes = None,
        mime_type: str = "application/pdf"
    ) -> Dict[str, Any]:
        """
        문서 처리 및 텍스트 추출
        
        Args:
            file_path: 파일 경로 (file_content와 둘 중 하나 필수)
            file_content: 파일 바이너리 내용
            mime_type: MIME 타입 (application/pdf, image/png 등)
        
        Returns:
            {
                "text": "추출된 전체 텍스트",
                "pages": [{"page_number": 1, "text": "...", "confidence": 0.98}],
                "entities": [...],
                "language": "ko",
                "confidence": 0.97
            }
        """
        if not self.is_available:
            # Document AI 사용 불가 시 fallback
            return await self._fallback_extract(file_path, file_content, mime_type)
        
        # 파일 로드
        if file_content is None and file_path:
            with open(file_path, "rb") as f:
                file_content = f.read()
        
        if file_content is None:
            raise ValueError("file_path 또는 file_content가 필요합니다")
        
        # Document AI 요청 생성
        client = self._get_client()
        
        raw_document = documentai.RawDocument(
            content=file_content,
            mime_type=mime_type
        )
        
        request = documentai.ProcessRequest(
            name=self._get_processor_name(),
            raw_document=raw_document
        )
        
        # 동기 처리 (작은 문서용)
        # 큰 문서는 batch_process_documents 사용 권장
        try:
            result = client.process_document(request=request)
            document = result.document
            
            # 결과 파싱
            pages = []
            for page in document.pages:
                page_text = self._extract_page_text(page, document.text)
                pages.append({
                    "page_number": page.page_number,
                    "text": page_text,
                    "confidence": page.layout.confidence if page.layout else 0.0,
                    "width": page.dimension.width if page.dimension else 0,
                    "height": page.dimension.height if page.dimension else 0
                })
            
            return {
                "text": document.text,
                "pages": pages,
                "entities": self._extract_entities(document),
                "language": document.detected_languages[0].language_code if document.detected_languages else "unknown",
                "confidence": self._calculate_avg_confidence(document),
                "source": "document_ai"
            }
        
        except Exception as e:
            # Document AI 실패 시 fallback
            print(f"Document AI 처리 실패: {e}, fallback 사용")
            return await self._fallback_extract(file_path, file_content, mime_type)
    
    def _extract_page_text(self, page, full_text: str) -> str:
        """페이지별 텍스트 추출"""
        if not page.layout or not page.layout.text_anchor:
            return ""
        
        text_segments = []
        for segment in page.layout.text_anchor.text_segments:
            start = int(segment.start_index) if segment.start_index else 0
            end = int(segment.end_index) if segment.end_index else len(full_text)
            text_segments.append(full_text[start:end])
        
        return "".join(text_segments)
    
    def _extract_entities(self, document) -> list:
        """엔티티 추출 (이름, 날짜 등)"""
        entities = []
        for entity in document.entities:
            entities.append({
                "type": entity.type_,
                "text": entity.mention_text,
                "confidence": entity.confidence
            })
        return entities
    
    def _calculate_avg_confidence(self, document) -> float:
        """평균 신뢰도 계산"""
        if not document.pages:
            return 0.0
        
        total = sum(
            page.layout.confidence if page.layout else 0.0
            for page in document.pages
        )
        return total / len(document.pages)
    
    async def _fallback_extract(
        self,
        file_path: str = None,
        file_content: bytes = None,
        mime_type: str = "application/pdf"
    ) -> Dict[str, Any]:
        """
        Document AI 사용 불가 시 로컬 라이브러리로 fallback
        
        pypdf, python-docx 등 사용
        """
        text = ""
        
        if file_path:
            ext = Path(file_path).suffix.lower()
        else:
            ext = self._guess_ext_from_mime(mime_type)
        
        try:
            if ext == ".pdf":
                text = self._extract_pdf(file_path, file_content)
            elif ext == ".docx":
                text = self._extract_docx(file_path, file_content)
            elif ext in [".txt", ".md"]:
                if file_content:
                    text = file_content.decode("utf-8", errors="ignore")
                elif file_path:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
            else:
                text = f"[지원하지 않는 형식: {ext}]"
        
        except Exception as e:
            text = f"[추출 실패: {str(e)}]"
        
        return {
            "text": text,
            "pages": [{"page_number": 1, "text": text, "confidence": 0.0}],
            "entities": [],
            "language": "ko",
            "confidence": 0.0,
            "source": "fallback"
        }
    
    def _extract_pdf(self, file_path: str = None, file_content: bytes = None) -> str:
        """PDF 텍스트 추출 (pypdf)"""
        try:
            from pypdf import PdfReader
            import io
            
            if file_content:
                reader = PdfReader(io.BytesIO(file_content))
            else:
                reader = PdfReader(file_path)
            
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
            
            return "\n\n".join(texts)
        
        except ImportError:
            return "[pypdf 패키지가 필요합니다]"
    
    def _extract_docx(self, file_path: str = None, file_content: bytes = None) -> str:
        """DOCX 텍스트 추출 (python-docx)"""
        try:
            from docx import Document
            import io
            
            if file_content:
                doc = Document(io.BytesIO(file_content))
            else:
                doc = Document(file_path)
            
            texts = [para.text for para in doc.paragraphs]
            return "\n".join(texts)
        
        except ImportError:
            return "[python-docx 패키지가 필요합니다]"
    
    def _guess_ext_from_mime(self, mime_type: str) -> str:
        """MIME 타입에서 확장자 추측"""
        mapping = {
            "application/pdf": ".pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "text/plain": ".txt",
            "text/markdown": ".md",
            "image/png": ".png",
            "image/jpeg": ".jpg"
        }
        return mapping.get(mime_type, ".txt")


# 싱글톤 인스턴스
_document_ai_service = None


def get_document_ai_service() -> DocumentAIService:
    """Document AI 서비스 싱글톤 반환"""
    global _document_ai_service
    if _document_ai_service is None:
        _document_ai_service = DocumentAIService()
    return _document_ai_service
